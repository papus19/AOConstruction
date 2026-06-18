"""
exports.py — Génération des fichiers téléchargeables et aperçus HTML
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Responsabilités :
  - Export Word (.docx)
  - Export Excel (.xlsx)
  - Export PDF (ReportLab)
  - Aperçus HTML (Word et Excel) pour affichage dans Streamlit
  - Affichage des boutons téléchargement + aperçus (utilise st.)

Ce fichier contient la seule zone où Streamlit est autorisé dans ce module.
Les fonctions _generer_* sont pures (retournent des bytes).
La fonction afficher_telechargements() utilise st.
"""

import io
import json
from datetime import datetime

import streamlit as st
import streamlit.components.v1 as components


# ─────────────────────────────────────────────────────────────────────────────
# 1. EXPORT WORD (.docx)
# ─────────────────────────────────────────────────────────────────────────────

def generer_docx(
    technique: dict,
    soumission: dict,
    analyse: dict,
    user: dict,
) -> bytes | None:
    """
    Génère le document Word complet (mémoire technique + devis détaillé).

    Args:
        technique: Mémoire technique généré par offre_technique.py.
        soumission: Données financières calculées par chiffrage.py.
        analyse: Dict d'analyse du projet.
        user: Profil entreprise.

    Returns:
        bytes du fichier .docx, None si erreur.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        BF = RGBColor(0x1E, 0x3A, 0x5F)
        BM = RGBColor(0x2E, 0x75, 0xB6)

        doc = Document()
        for s in doc.sections:
            s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(2.5)

        def h1(txt):
            p = doc.add_heading(txt, level=1)
            for r in p.runs:
                r.font.color.rgb = BF
            return p

        def h2(txt):
            p = doc.add_heading(txt, level=2)
            for r in p.runs:
                r.font.color.rgb = BM
            return p

        def body(txt, bold=False):
            if not txt:
                return
            p = doc.add_paragraph()
            r = p.add_run(str(txt))
            r.bold = bold
            r.font.size = Pt(11)

        def bullet(txt):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(str(txt)).font.size = Pt(11)

        # Page de titre
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(48)
        r = p.add_run(technique.get("titre_offre") or technique.get("titre", "Offre de services"))
        r.bold = True
        r.font.size = Pt(20)
        r.font.color.rgb = BF

        nom_projet = analyse.get("nom_projet") or (analyse.get("exigences") or {}).get("nom_projet", "")
        client     = analyse.get("client")     or (analyse.get("exigences") or {}).get("client", "")

        for txt, sz in [
            (user.get("nom_entreprise", ""),             14),
            (f"RBQ : {user.get('licence_rbq', 'N/A')}", 11),
            (f"Date : {datetime.today().strftime('%d %B %Y')}", 11),
            (f"Projet : {nom_projet}",                   12),
            (f"Client : {client}",                       11),
        ]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(txt)
            r.font.size = Pt(sz)
            r.font.color.rgb = BM if sz >= 12 else RGBColor(0x44, 0x44, 0x44)

        doc.add_page_break()

        # Corps du mémoire
        h1("1. Introduction")
        body(technique.get("introduction", ""))

        h1("2. Compréhension du mandat")
        body(technique.get("comprehension_projet") or technique.get("comprehension_mandat", ""))

        h1("3. Méthodologie")
        phases = (technique.get("approche_methodologique") or {}).get("phases") or technique.get("methodologie", [])
        for phase in phases:
            h2(f"{phase.get('nom') or phase.get('phase', '')}  ({phase.get('duree', '')})")
            body(phase.get("description", ""))

        h1("4. Approche technique par poste")
        for pt in technique.get("postes_techniques", []):
            h2(pt.get("titre", ""))
            body(pt.get("approche", ""))
            for n in pt.get("normes_applicables", []):
                bullet(n)

        h1("5. Équipe")
        equipe = technique.get("equipe_proposee") or technique.get("equipe", [])
        for m in equipe:
            h2(f"{m.get('role', '')} — {m.get('nom', '')}")
            body(m.get("experience") or m.get("qualification", ""))

        h1("6. Garanties")
        garanties = technique.get("garanties_qualite") or technique.get("garanties", [])
        for g in garanties:
            bullet(g)

        # Sections supplémentaires
        for s in technique.get("sections_supplementaires", []):
            h1(s.get("titre", "Section complémentaire"))
            body(s.get("contenu", ""))

        doc.add_page_break()

        # Bordereau de prix
        h1("7. Bordereau de prix ventilés")
        body(
            "Les prix ci-dessous sont présentés conformément au bordereau de soumission "
            "fourni par le donneur d'ouvrage. Tous les prix sont en dollars canadiens, "
            "excluant les taxes."
        )
        doc.add_paragraph()

        sections_bord = st.session_state.get("offre_data", {}).get("bordereau_ao", [])
        if sections_bord:
            grand_total = 0.0
            for sec in sections_bord:
                # Titre de section en bandeau
                p_sec = doc.add_paragraph()
                p_sec.paragraph_format.space_before = Pt(10)
                r_sec = p_sec.add_run(sec.get("titre", "").upper())
                r_sec.bold = True
                r_sec.font.size = Pt(10)
                r_sec.font.color.rgb = BF

                # Tableau 2 colonnes — sans Source
                tbl_sec = doc.add_table(rows=1, cols=2)
                tbl_sec.style = "Table Grid"
                hdr_cells = tbl_sec.rows[0].cells
                for i, (txt, w) in enumerate([
                    ("DESCRIPTION", Cm(13.5)),
                    ("PRIX SOUMIS ($)", Cm(4)),
                ]):
                    hdr_cells[i].text = txt
                    hdr_cells[i].width = w
                    for run in hdr_cells[i].paragraphs[0].runs:
                        run.bold = True
                        run.font.size = Pt(9)

                sous_total_sec = 0.0
                for it in sec.get("items", []):
                    row  = tbl_sec.add_row().cells
                    prix = float(it.get("prix_unitaire") or 0)
                    qte  = float(it.get("quantite") or 1)
                    total_item = round(prix * qte, 2)
                    sous_total_sec += total_item
                    row[0].text = it.get("description", "")
                    row[1].text = f"{total_item:,.2f} $" if total_item else "À compléter"
                    row[0].width = Cm(13.5)
                    row[1].width = Cm(4)
                    for cell in row:
                        for run in cell.paragraphs[0].runs:
                            run.font.size = Pt(9)

                # Ligne sous-total de section
                row_st = tbl_sec.add_row().cells
                row_st[0].text = f"Sous-total — {sec.get('titre','')}"
                row_st[1].text = f"{sous_total_sec:,.2f} $"
                row_st[0].width = Cm(13.5)
                row_st[1].width = Cm(4)
                for cell in row_st:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                        run.font.size = Pt(9)
                        run.font.color.rgb = BF
                grand_total += sous_total_sec
                doc.add_paragraph()

            # Tableau totaux + taxes
            doc.add_paragraph()
            contig_pct = (st.session_state.get("offre_data") or {}).get("contingence", 10) or 10
            contig_v   = round(grand_total * contig_pct / 100, 2)
            avant_tx_v = round(grand_total + contig_v, 2)
            tps_v      = round(avant_tx_v * 0.05, 2)
            tvq_v      = round(avant_tx_v * 0.09975, 2)
            ttc_v      = round(avant_tx_v + tps_v + tvq_v, 2)

            lignes_tot = [
                ("Sous-total H.T.",              f"{grand_total:,.2f} $",  False),
                (f"Contingence ({contig_pct}%)", f"{contig_v:,.2f} $",     False),
                ("Total avant taxes",             f"{avant_tx_v:,.2f} $",  False),
                ("TPS (5%)",                      f"{tps_v:,.2f} $",        False),
                ("TVQ (9,975%)",                  f"{tvq_v:,.2f} $",        False),
                ("TOTAL TTC",                     f"{ttc_v:,.2f} $",        True),
            ]
            tbl_tot = doc.add_table(rows=len(lignes_tot), cols=2)
            tbl_tot.style = "Table Grid"
            for idx, (lbl, val, bold) in enumerate(lignes_tot):
                cells = tbl_tot.rows[idx].cells
                cells[0].text = lbl
                cells[1].text = val
                cells[0].width = Cm(13.5)
                cells[1].width = Cm(4)
                for cell in cells:
                    for run in cell.paragraphs[0].runs:
                        run.bold = bold
                        run.font.size = Pt(10 if bold else 9)
                        if bold:
                            run.font.color.rgb = BF

            # Bloc signature
            doc.add_paragraph()
            doc.add_paragraph()
            p_sig_titre = doc.add_paragraph()
            r_sig_titre = p_sig_titre.add_run("SIGNATURE DU SOUMISSIONNAIRE")
            r_sig_titre.bold = True
            r_sig_titre.font.size = Pt(11)
            r_sig_titre.font.color.rgb = BF

            tbl_sig = doc.add_table(rows=5, cols=2)
            tbl_sig.style = "Table Grid"
            sig_data = [
                ("Raison sociale",         user.get("nom_entreprise", "")),
                ("Licence RBQ",            user.get("licence_rbq", "")),
                ("Nom (lettres moulées)",  user.get("contact_nom", "").upper()),
                ("Titre / Fonction",       user.get("contact_titre", "")),
                ("Signature / Date",       "________________________________  /  ___________"),
            ]
            for idx, (lbl, val) in enumerate(sig_data):
                cells = tbl_sig.rows[idx].cells
                cells[0].text = lbl
                cells[1].text = val
                cells[0].width = Cm(5)
                cells[1].width = Cm(12.5)
                for run in cells[0].paragraphs[0].runs:
                    run.bold = True
                    run.font.size = Pt(9)
                for run in cells[1].paragraphs[0].runs:
                    run.font.size = Pt(9)
        else:
            body("(Aucun bordereau de prix chargé — veuillez compléter l'onglet Offre Financière)")

        # ── Aucun saut de page supplémentaire — le bordereau est déjà la section 7
        # Bloc info entreprise en pied du bordereau
        doc.add_paragraph()
        tbl_info = doc.add_table(rows=1, cols=2)
        tbl_info.style = "Table Grid"
        info_data = [
            f"{user.get('nom_entreprise', '')}",
            f"RBQ : {user.get('licence_rbq', '')}  |  {user.get('contact_telephone', '')}  |  {user.get('contact_email', '')}",
        ]
        for col_idx, txt in enumerate(info_data):
            c = tbl_info.rows[0].cells[col_idx]
            c.text = txt
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(9)
                run.bold = (col_idx == 0)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    except Exception as e:
        st.error(f"❌ Erreur Word : {e}")
        import traceback
        st.code(traceback.format_exc())
        return None


# ─────────────────────────────────────────────────────────────────────────────
# 2. EXPORT EXCEL (.xlsx)
# ─────────────────────────────────────────────────────────────────────────────

def generer_xlsx(soumission: dict, analyse: dict, user: dict) -> bytes | None:
    """
    Génère le tableur Excel du bordereau de prix par section.
    - Sections avec bandeau sous-titre
    - Sans colonne Source
    - Colonnes adaptées au contenu
    - Bloc signature en bas

    Args:
        soumission: Données financières calculées.
        analyse: Dict d'analyse du projet.
        user: Profil entreprise.

    Returns:
        bytes du fichier .xlsx, None si erreur.
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter

        BF = "1E3A5F"
        BM = "2E75B6"
        BC = "D5E8F0"
        GC = "F5F7FA"
        BL = "FFFFFF"
        THIN = Side(style="thin", color="CCCCCC")
        brd  = Border(top=THIN, bottom=THIN, left=THIN, right=THIN)

        def cs(ws, row, col, val, bold=False, bg=None, fg="000000",
               align="left", sz=10, nfmt=None, wrap=False):
            c = ws.cell(row=row, column=col, value=val)
            c.font = Font(name="Arial", bold=bold, color=fg, size=sz)
            if bg:
                c.fill = PatternFill("solid", fgColor=bg)
            c.alignment = Alignment(horizontal=align, vertical="center", wrap_text=wrap)
            if nfmt:
                c.number_format = nfmt
            c.border = brd
            return c

        def merge_row(ws, row, nb_cols, val, bold=False, bg=BF, fg=BL, sz=10):
            last_col = get_column_letter(nb_cols)
            ws.merge_cells(f"A{row}:{last_col}{row}")
            c = ws.cell(row, 1, val)
            c.font      = Font(name="Arial", bold=bold, color=fg, size=sz)
            c.fill      = PatternFill("solid", fgColor=bg)
            c.alignment = Alignment(horizontal="left", vertical="center")
            c.border    = brd
            ws.row_dimensions[row].height = 18

        wb  = Workbook()
        ws1 = wb.active
        ws1.title = "Bordereau de prix"

        # Largeurs des colonnes : Description large, Qté, Unité, Total — pas de Source
        # Col: A=No  B=Description  C=Qté  D=Unité  E=Total ($)
        NB_COLS = 5
        col_widths = {"A": 6, "B": 60, "C": 8, "D": 12, "E": 16}
        for col_ltr, w in col_widths.items():
            ws1.column_dimensions[col_ltr].width = w

        nom_projet = analyse.get("nom_projet") or (analyse.get("exigences") or {}).get("nom_projet", "")
        client     = analyse.get("client")     or (analyse.get("exigences") or {}).get("client", "")
        contact    = user.get("contact_nom", "")
        rbq        = user.get("licence_rbq", "")

        r = 1
        # En-tête principal
        ws1.merge_cells(f"A{r}:E{r}")
        c = ws1.cell(r, 1, f"BORDEREAU DE PRIX — {user.get('nom_entreprise','').upper()}")
        c.font = Font(name="Arial", bold=True, size=14, color=BL)
        c.fill = PatternFill("solid", fgColor=BF)
        c.alignment = Alignment(horizontal="center", vertical="center")
        ws1.row_dimensions[r].height = 28
        r += 1

        ws1.merge_cells(f"A{r}:E{r}")
        c = ws1.cell(r, 1, f"Projet : {nom_projet}  |  Client : {client}  |  {datetime.today().strftime('%d/%m/%Y')}")
        c.font = Font(name="Arial", size=10, color=BL)
        c.fill = PatternFill("solid", fgColor=BM)
        c.alignment = Alignment(horizontal="center", vertical="center")
        ws1.row_dimensions[r].height = 18
        r += 2

        # En-têtes colonnes
        entetes = ["No", "Description", "Qté", "Unité", "Total ($)"]
        for col, h in enumerate(entetes, 1):
            cs(ws1, r, col, h, bold=True, bg=BF, fg=BL, align="center", sz=9)
        ws1.row_dimensions[r].height = 20
        r += 1

        # Données par section
        sections_bord = st.session_state.get("offre_data", {}).get("bordereau_ao", [])
        grand_total = 0.0

        if sections_bord:
            for sec in sections_bord:
                # Bandeau section
                merge_row(ws1, r, NB_COLS, f"  {sec.get('titre','').upper()}",
                          bold=True, bg=BM, fg=BL, sz=9)
                r += 1

                sous_total_sec = 0.0
                for i, it in enumerate(sec.get("items", [])):
                    bg     = GC if i % 2 == 0 else BL
                    prix   = float(it.get("prix_unitaire") or 0)
                    qte    = float(it.get("quantite") or 1)
                    total  = round(prix * qte, 2)
                    sous_total_sec += total

                    cs(ws1, r, 1, it.get("no",""),          bg=bg, sz=9, align="center")
                    cs(ws1, r, 2, it.get("description",""), bg=bg, sz=9, wrap=True)
                    cs(ws1, r, 3, qte,                      bg=bg, sz=9, align="center")
                    cs(ws1, r, 4, it.get("unite","forfait"),bg=bg, sz=9)
                    cs(ws1, r, 5, total if total else None,  bg=bg, sz=9, align="right",
                       nfmt='"$"#,##0.00')
                    ws1.row_dimensions[r].height = 16
                    r += 1

                grand_total += sous_total_sec

                # Ligne sous-total section
                merge_row(ws1, r, NB_COLS - 1,
                          f"  Sous-total — {sec.get('titre','')}",
                          bold=True, bg=BC, fg=BF, sz=9)
                c_st = ws1.cell(r, NB_COLS, sous_total_sec)
                c_st.font         = Font(name="Arial", bold=True, size=9, color=BF)
                c_st.fill         = PatternFill("solid", fgColor=BC)
                c_st.number_format = '"$"#,##0.00'
                c_st.alignment    = Alignment(horizontal="right", vertical="center")
                c_st.border       = brd
                ws1.row_dimensions[r].height = 16
                r += 1

            r += 1
        else:
            # Fallback postes plats si pas de bordereau structuré
            postes = soumission.get("postes") or []
            for i, p in enumerate(postes):
                bg    = GC if i % 2 == 0 else BL
                total = float(p.get("total_poste", p.get("total", 0)))
                grand_total += total
                cs(ws1, r, 1, p.get("code",""),          bg=bg, sz=9)
                cs(ws1, r, 2, p.get("description",""),   bg=bg, sz=9, wrap=True)
                cs(ws1, r, 3, p.get("quantite",1),       bg=bg, sz=9, align="center")
                cs(ws1, r, 4, p.get("unite","forfait"),  bg=bg, sz=9)
                cs(ws1, r, 5, total,                      bg=bg, sz=9, align="right", nfmt='"$"#,##0.00')
                ws1.row_dimensions[r].height = 16
                r += 1
            r += 1

        # Totaux + taxes
        contig_pct = (st.session_state.get("offre_data") or {}).get("contingence", 10) or 10
        contig_v   = round(grand_total * contig_pct / 100, 2)
        avant_tx_v = round(grand_total + contig_v, 2)
        tps_v      = round(avant_tx_v * 0.05, 2)
        tvq_v      = round(avant_tx_v * 0.09975, 2)
        ttc_v      = round(avant_tx_v + tps_v + tvq_v, 2)

        def ltotal(lbl, val, row, bold=False, bg=BC, fg=BF):
            ws1.merge_cells(f"A{row}:D{row}")
            c1 = ws1.cell(row, 1, lbl)
            c1.font      = Font(name="Arial", bold=bold, size=10 if bold else 9, color=fg)
            c1.fill      = PatternFill("solid", fgColor=bg)
            c1.alignment = Alignment(horizontal="right", vertical="center")
            c1.border    = brd
            c2 = ws1.cell(row, 5, val)
            c2.font         = Font(name="Arial", bold=bold, size=10 if bold else 9, color=fg)
            c2.fill         = PatternFill("solid", fgColor=bg)
            c2.number_format = '"$"#,##0.00'
            c2.alignment    = Alignment(horizontal="right")
            c2.border       = brd
            ws1.row_dimensions[row].height = 18 if bold else 16

        ltotal("Sous-total H.T.",              grand_total,  r);           r += 1
        ltotal(f"Contingence ({contig_pct}%)", contig_v,     r);           r += 1
        ltotal("Total avant taxes",             avant_tx_v,   r);           r += 1
        ltotal("TPS (5%)",                      tps_v,        r);           r += 1
        ltotal("TVQ (9,975%)",                  tvq_v,        r);           r += 1
        ltotal("TOTAL TTC",                     ttc_v,        r,
               bold=True, bg=BF, fg=BL);                                    r += 2

        # Bloc signature
        merge_row(ws1, r, NB_COLS, "  SIGNATURE DU SOUMISSIONNAIRE",
                  bold=True, bg=BF, fg=BL, sz=10)
        ws1.row_dimensions[r].height = 20
        r += 1

        sig_lignes = [
            ("Raison sociale",         user.get("nom_entreprise", "")),
            ("Licence RBQ",            rbq),
            ("Nom (lettres moulées)",  contact.upper()),
            ("Titre / Fonction",       user.get("contact_titre", "")),
            ("Signature",              ""),
            ("Date",                   ""),
        ]
        for lbl, val in sig_lignes:
            cs(ws1, r, 1, lbl,  bold=True, bg=GC, fg=BF, sz=9, align="right")
            ws1.merge_cells(f"B{r}:E{r}")
            c_val = ws1.cell(r, 2, val)
            c_val.font      = Font(name="Arial", size=9)
            c_val.fill      = PatternFill("solid", fgColor=BL)
            c_val.border    = brd
            c_val.alignment = Alignment(horizontal="left", vertical="center")
            ws1.row_dimensions[r].height = 18
            r += 1

        out = io.BytesIO()
        wb.save(out)
        return out.getvalue()

    except Exception as e:
        st.error(f"❌ Erreur Excel : {e}")
        import traceback
        st.code(traceback.format_exc())
        return None


# ─────────────────────────────────────────────────────────────────────────────
# 3. EXPORT PDF (ReportLab)
# ─────────────────────────────────────────────────────────────────────────────

def generer_pdf(
    technique: dict,
    soumission: dict,
    analyse: dict,
    user: dict,
) -> bytes | None:
    """
    Génère le PDF complet (mémoire technique + devis).

    Args:
        technique: Mémoire technique.
        soumission: Données financières.
        analyse: Dict d'analyse du projet.
        user: Profil entreprise.

    Returns:
        bytes du fichier .pdf, None si erreur.
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib.colors import HexColor, white, black
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer,
            Table, TableStyle, PageBreak, HRFlowable
        )
        from reportlab.lib.enums import TA_CENTER

        BF = HexColor("#1E3A5F")
        BM = HexColor("#2E75B6")
        BC = HexColor("#D5E8F0")
        GC = HexColor("#F5F7FA")

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf, pagesize=letter,
            topMargin=2.5*cm, bottomMargin=2.5*cm,
            leftMargin=2.5*cm, rightMargin=2.5*cm
        )

        S = {
            "titre":  ParagraphStyle("titre",  fontSize=20, textColor=BF,
                        spaceAfter=8, alignment=TA_CENTER, fontName="Helvetica-Bold"),
            "sous":   ParagraphStyle("sous",   fontSize=12, textColor=BM,
                        spaceAfter=4, alignment=TA_CENTER, fontName="Helvetica"),
            "info":   ParagraphStyle("info",   fontSize=10, textColor=HexColor("#444444"),
                        spaceAfter=3, alignment=TA_CENTER, fontName="Helvetica"),
            "body":   ParagraphStyle("body",   fontSize=10, textColor=black,
                        spaceAfter=6, leading=14, fontName="Helvetica"),
            "h2":     ParagraphStyle("h2",     fontSize=11, textColor=BF,
                        spaceBefore=8, spaceAfter=3, fontName="Helvetica-Bold"),
            "bullet": ParagraphStyle("bullet", fontSize=10, textColor=black,
                        spaceAfter=3, leading=13, fontName="Helvetica", leftIndent=15),
            "small":  ParagraphStyle("small",  fontSize=8,  textColor=HexColor("#666666"),
                        spaceAfter=2, fontName="Helvetica-Oblique"),
        }

        story = []

        def h1(txt):
            t = Table(
                [[Paragraph(txt, ParagraphStyle("hh", fontSize=11,
                    textColor=white, fontName="Helvetica-Bold"))]],
                colWidths=[18.5*cm]
            )
            t.setStyle(TableStyle([
                ("BACKGROUND",    (0,0), (-1,-1), BF),
                ("TOPPADDING",    (0,0), (-1,-1), 5),
                ("BOTTOMPADDING", (0,0), (-1,-1), 5),
                ("LEFTPADDING",   (0,0), (-1,-1), 8),
            ]))
            story.append(Spacer(1, 0.3*cm))
            story.append(t)
            story.append(Spacer(1, 0.2*cm))

        nom_projet = analyse.get("nom_projet") or (analyse.get("exigences") or {}).get("nom_projet", "")
        titre_off  = technique.get("titre_offre") or technique.get("titre", "Offre de services")

        story += [
            Spacer(1, 2*cm),
            Paragraph(titre_off, S["titre"]),
            Spacer(1, 0.3*cm),
            Paragraph(user.get("nom_entreprise", ""), S["sous"]),
            Paragraph(f"RBQ : {user.get('licence_rbq', '')}", S["info"]),
            Paragraph(f"Date : {datetime.today().strftime('%d %B %Y')}", S["info"]),
            Paragraph(f"Projet : {nom_projet}", S["info"]),
            Spacer(1, 0.5*cm),
            HRFlowable(width="100%", thickness=2, color=BM),
            PageBreak(),
        ]

        h1("1. Introduction")
        story.append(Paragraph(technique.get("introduction", ""), S["body"]))

        h1("2. Compréhension du mandat")
        story.append(Paragraph(
            technique.get("comprehension_projet") or technique.get("comprehension_mandat", ""),
            S["body"]
        ))

        h1("3. Méthodologie")
        phases = (technique.get("approche_methodologique") or {}).get("phases") or technique.get("methodologie", [])
        for ph in phases:
            story.append(Paragraph(
                f"{ph.get('nom') or ph.get('phase','')} ({ph.get('duree','')})",
                S["h2"]
            ))
            story.append(Paragraph(ph.get("description", ""), S["body"]))

        h1("4. Approche technique par poste")
        for pt in technique.get("postes_techniques", []):
            story.append(Paragraph(pt.get("titre", ""), S["h2"]))
            story.append(Paragraph(pt.get("approche", ""), S["body"]))
            for n in pt.get("normes_applicables", []):
                story.append(Paragraph(f"• {n}", S["bullet"]))

        h1("5. Garanties")
        garanties = technique.get("garanties_qualite") or technique.get("garanties", [])
        for g in garanties:
            story.append(Paragraph(f"• {g}", S["bullet"]))

        story.append(PageBreak())
        h1("6. Bordereau de prix ventilés")

        sections_bord_pdf = st.session_state.get("offre_data", {}).get("bordereau_ao", [])
        grand_total_pdf   = 0.0

        if sections_bord_pdf:
            for sec in sections_bord_pdf:
                # Bandeau section
                t_sec = Table(
                    [[Paragraph(sec.get("titre","").upper(),
                        ParagraphStyle("sh", fontSize=9, textColor=white,
                            fontName="Helvetica-Bold"))]],
                    colWidths=[18.5*cm]
                )
                t_sec.setStyle(TableStyle([
                    ("BACKGROUND",    (0,0), (-1,-1), BM),
                    ("TOPPADDING",    (0,0), (-1,-1), 4),
                    ("BOTTOMPADDING", (0,0), (-1,-1), 4),
                    ("LEFTPADDING",   (0,0), (-1,-1), 8),
                ]))
                story.append(Spacer(1, 0.25*cm))
                story.append(t_sec)

                # En-tête colonnes
                entete_sec = [["Description", "Qté", "Unité", "Prix unit. $", "Total $"]]
                rows_sec   = []
                sous_total  = 0.0

                for it in sec.get("items", []):
                    prix  = float(it.get("prix_unitaire") or 0)
                    qte   = float(it.get("quantite") or 1)
                    total = round(prix * qte, 2)
                    sous_total += total
                    rows_sec.append([
                        Paragraph(it.get("description",""), ParagraphStyle(
                            "d", fontSize=8, fontName="Helvetica", leading=10)),
                        str(int(qte)),
                        it.get("unite","forfait"),
                        f"${prix:,.2f}" if prix else "À compléter",
                        f"${total:,.2f}" if total else "",
                    ])

                # Ligne sous-total section
                rows_sec.append([
                    Paragraph(f"<b>Sous-total — {sec.get('titre','')}</b>",
                        ParagraphStyle("st", fontSize=8, fontName="Helvetica-Bold",
                            textColor=HexColor("#1E3A5F"))),
                    "", "", "",
                    Paragraph(f"<b>${sous_total:,.2f}</b>",
                        ParagraphStyle("stv", fontSize=8, fontName="Helvetica-Bold",
                            textColor=HexColor("#1E3A5F"))),
                ])
                grand_total_pdf += sous_total

                t_items = Table(
                    entete_sec + rows_sec,
                    colWidths=[10*cm, 1.2*cm, 2*cm, 2.8*cm, 2.5*cm]
                )
                n_rows = len(rows_sec)
                t_items.setStyle(TableStyle([
                    ("BACKGROUND",    (0,0), (-1,0),  BF),
                    ("TEXTCOLOR",     (0,0), (-1,0),  white),
                    ("FONTNAME",      (0,0), (-1,0),  "Helvetica-Bold"),
                    ("FONTSIZE",      (0,0), (-1,-1), 8),
                    ("GRID",          (0,0), (-1,-1), 0.3, HexColor("#CCCCCC")),
                    ("ALIGN",         (1,0), (-1,-1), "RIGHT"),
                    ("VALIGN",        (0,0), (-1,-1), "TOP"),
                    ("TOPPADDING",    (0,0), (-1,-1), 3),
                    ("BOTTOMPADDING", (0,0), (-1,-1), 3),
                    ("BACKGROUND",    (0, n_rows), (-1, n_rows), HexColor("#D5E8F0")),
                    ("ROWBACKGROUNDS",(0,1), (-1, n_rows-1), [GC, white]),
                ]))
                story.append(t_items)
        else:
            story.append(Paragraph(
                "(Aucun bordereau chargé — complétez l'onglet Offre Financière)",
                S["body"]
            ))

        # Tableau totaux + taxes
        story.append(Spacer(1, 0.5*cm))
        contig_pct_pdf = (st.session_state.get("offre_data") or {}).get("contingence", 10) or 10
        contig_pdf     = round(grand_total_pdf * contig_pct_pdf / 100, 2)
        avant_tx_pdf   = round(grand_total_pdf + contig_pdf, 2)
        tps_pdf        = round(avant_tx_pdf * 0.05, 2)
        tvq_pdf        = round(avant_tx_pdf * 0.09975, 2)
        ttc_pdf        = round(avant_tx_pdf + tps_pdf + tvq_pdf, 2)

        totaux = [
            ["Sous-total H.T.",              f"${grand_total_pdf:,.2f}"],
            [f"Contingence ({contig_pct_pdf}%)", f"${contig_pdf:,.2f}"],
            ["Total avant taxes",             f"${avant_tx_pdf:,.2f}"],
            ["TPS (5%)",                      f"${tps_pdf:,.2f}"],
            ["TVQ (9,975%)",                  f"${tvq_pdf:,.2f}"],
            ["TOTAL TTC",                     f"${ttc_pdf:,.2f}"],
        ]
        t_tot = Table(totaux, colWidths=[14*cm, 4.5*cm])
        t_tot.setStyle(TableStyle([
            ("FONTNAME",      (0,0), (-1,-2), "Helvetica"),
            ("FONTNAME",      (0,-1), (-1,-1), "Helvetica-Bold"),
            ("FONTSIZE",      (0,0), (-1,-1), 9),
            ("ALIGN",         (0,0), (-1,-1), "RIGHT"),
            ("LINEABOVE",     (0,-3), (-1,-3), 1, BM),
            ("BACKGROUND",    (0,-1), (-1,-1), BF),
            ("TEXTCOLOR",     (0,-1), (-1,-1), white),
            ("TOPPADDING",    (0,0), (-1,-1), 3),
            ("BOTTOMPADDING", (0,0), (-1,-1), 3),
        ]))
        story.append(t_tot)

        # Bloc signature
        story.append(Spacer(1, 1*cm))
        story.append(PageBreak())
        h1("7. Signature du soumissionnaire")
        story.append(Spacer(1, 0.4*cm))

        sig_rows = [
            ["Raison sociale",        user.get("nom_entreprise", "")],
            ["Licence RBQ",           user.get("licence_rbq", "")],
            ["Téléphone",             user.get("contact_telephone", "")],
            ["Courriel",              user.get("contact_email", "")],
            ["Nom (lettres moulées)", user.get("contact_nom", "").upper()],
            ["Titre / Fonction",      user.get("contact_titre", "")],
            ["Signature",             ""],
            ["Date",                  ""],
        ]
        t_sig = Table(sig_rows, colWidths=[5*cm, 13.5*cm])
        t_sig.setStyle(TableStyle([
            ("FONTNAME",      (0,0), (0,-1), "Helvetica-Bold"),
            ("FONTNAME",      (1,0), (1,-1), "Helvetica"),
            ("FONTSIZE",      (0,0), (-1,-1), 9),
            ("GRID",          (0,0), (-1,-1), 0.3, HexColor("#CCCCCC")),
            ("BACKGROUND",    (0,0), (0,-1), GC),
            ("TEXTCOLOR",     (0,0), (0,-1), BF),
            ("ALIGN",         (0,0), (0,-1), "RIGHT"),
            ("TOPPADDING",    (0,0), (-1,-1), 5),
            ("BOTTOMPADDING", (0,0), (-1,-1), 5),
            ("ROWHEIGHT",     (0,-2), (-1,-1), 1.2*cm),   # Signature et Date plus hautes
        ]))
        story.append(t_sig)

        doc.build(story)
        return buf.getvalue()

    except Exception as e:
        st.error(f"❌ Erreur PDF : {e}")
        import traceback
        st.code(traceback.format_exc())
        return None


# ─────────────────────────────────────────────────────────────────────────────
# 4. APERÇUS HTML (pour affichage Streamlit inline)
# ─────────────────────────────────────────────────────────────────────────────

def apercu_word_html(offre_tech: dict, exigences: dict, user: dict) -> str:
    """
    Construit un aperçu HTML du mémoire technique (rendu fidèle au Word).

    Args:
        offre_tech: Mémoire technique.
        exigences: Dict d'exigences du projet.
        user: Profil entreprise.

    Returns:
        Chaîne HTML complète prête pour st.markdown() ou components.html().
    """
    bleu_f = "#1E3A5F"
    bleu_m = "#2E75B6"
    gris   = "#666"

    def section(titre, contenu):
        return (
            f"<div style='margin-bottom:16px;'>"
            f"<div style='background:{bleu_f};color:white;font-weight:bold;"
            f"font-size:12px;padding:6px 12px;border-radius:3px;margin-bottom:8px;'>"
            f"{titre}</div>"
            f"<div style='font-size:11px;line-height:1.7;color:#212529;padding:0 6px;'>"
            f"{contenu}</div></div>"
        )

    parts = [
        f"<div style='text-align:center;padding:20px 0 16px;"
        f"border-bottom:3px solid {bleu_m};margin-bottom:20px;'>"
        f"<div style='font-size:19px;font-weight:bold;color:{bleu_f};'>"
        f"{offre_tech.get('titre_offre') or offre_tech.get('titre','Offre de services')}</div>"
        f"<div style='font-size:13px;color:{bleu_m};margin-top:5px;'>"
        f"{user.get('nom_entreprise','')}</div>"
        f"<div style='font-size:10px;color:{gris};margin-top:4px;'>"
        f"RBQ : {user.get('licence_rbq','')} &nbsp;|&nbsp; "
        f"Préparé pour : {exigences.get('client') or 'le donneur d\u2019ouvrage'}"
        f"</div></div>"
    ]

    parts.append(section("1. Introduction", offre_tech.get("introduction", "")))
    parts.append(section(
        "2. Compréhension du projet",
        offre_tech.get("comprehension_projet") or offre_tech.get("comprehension_mandat", "")
    ))

    approche    = offre_tech.get("approche_methodologique", {}) or {}
    phases      = approche.get("phases") or offre_tech.get("methodologie", [])
    phases_html = "".join(
        f"<div style='margin:6px 0 6px 12px;'>"
        f"<span style='font-weight:bold;color:{bleu_m};'>"
        f"{ph.get('nom') or ph.get('phase','')} ({ph.get('duree','')})</span>"
        f" — {ph.get('description','')}</div>"
        for ph in phases
    )
    parts.append(section(
        "3. Approche méthodologique",
        f"{approche.get('description','')}<br>{phases_html}"
    ))

    equipe  = offre_tech.get("equipe_proposee") or offre_tech.get("equipe", [])
    eq_html = "".join(
        f"<div style='margin:5px 0 5px 12px;'>"
        f"<span style='font-weight:bold;color:{bleu_m};'>"
        f"{m.get('role','')} : {m.get('nom','')}</span><br>"
        f"<span style='font-size:10px;color:{gris};'>"
        f"{m.get('experience','') or m.get('qualification','')}</span></div>"
        for m in equipe
    )
    parts.append(section("4. Équipe proposée", eq_html))

    garanties = offre_tech.get("garanties_qualite") or offre_tech.get("garanties", [])
    gar_html  = "".join(f"<div style='margin:3px 0 3px 12px;'>• {g}</div>" for g in garanties)
    parts.append(section("5. Garanties", gar_html))

    # Sections supplémentaires ajoutées manuellement
    for i, s in enumerate(offre_tech.get("sections_supplementaires", []), start=6):
        parts.append(section(
            f"{i}. {s.get('titre','')}",
            s.get("contenu","").replace("\n", "<br>")
        ))

    # Bordereau si disponible dans session
    sections_bord = st.session_state.get("offre_data", {}).get("bordereau_ao", [])
    if sections_bord:
        bord_html  = (
            "<p style='font-size:10px;color:#555;margin-bottom:10px;'>"
            "Prix en dollars canadiens, excluant les taxes applicables.</p>"
        )
        grand_total = 0.0
        for sec in sections_bord:
            rows_items     = ""
            sous_total_sec = 0.0
            for it in sec.get("items", []):
                prix        = float(it.get("prix_unitaire") or 0)
                qte         = float(it.get("quantite") or 1)
                total_item  = round(prix * qte, 2)
                sous_total_sec += total_item
                prix_txt    = f"{total_item:,.2f} $" if total_item else "<em style='color:#e74c3c;'>À compléter</em>"
                source_txt  = it.get("source_prix", "") or ""
                rows_items += (
                    f"<tr>"
                    f"<td style='padding:5px 8px;border:1px solid #ddd;'>- {it.get('description','')}</td>"
                    f"<td style='padding:5px 8px;border:1px solid #ddd;text-align:right;'>{prix_txt}</td>"
                    f"<td style='padding:5px 8px;border:1px solid #ddd;font-size:9px;color:{gris};'>{source_txt}</td>"
                    f"</tr>"
                )
            grand_total += sous_total_sec
            bord_html += (
                f"<div style='margin-bottom:12px;'>"
                f"<div style='font-weight:bold;font-size:11px;color:{bleu_f};padding:5px 8px;"
                f"background:#eef2f8;border-left:4px solid {bleu_m};margin-bottom:4px;'>"
                f"{sec.get('titre','')}</div>"
                f"<table style='width:100%;border-collapse:collapse;font-size:10px;'>"
                f"<thead><tr style='background:{bleu_f};color:white;'>"
                f"<th style='padding:4px 8px;text-align:left;'>Description</th>"
                f"<th style='padding:4px 8px;text-align:right;'>Prix soumis</th>"
                f"<th style='padding:4px 8px;text-align:left;'>Source</th>"
                f"</tr></thead><tbody>{rows_items}</tbody>"
                f"<tfoot><tr style='background:#f5f7fa;font-weight:bold;'>"
                f"<td style='padding:5px 8px;border:1px solid #ddd;'>Sous-total</td>"
                f"<td style='padding:5px 8px;border:1px solid #ddd;text-align:right;'>{sous_total_sec:,.2f} $</td>"
                f"<td style='padding:5px 8px;border:1px solid #ddd;'></td>"
                f"</tr></tfoot></table></div>"
            )
        bord_html += (
            f"<div style='margin-top:14px;padding:10px 14px;background:{bleu_f};"
            f"color:white;font-size:14px;font-weight:bold;border-radius:4px;text-align:right;'>"
            f"GRAND TOTAL (excluant les taxes) : {grand_total:,.2f} $</div>"
        )
        parts.append(section("Bordereau de prix ventilés", bord_html))

    return (
        f"<html><head><meta charset='utf-8'>"
        f"<style>body{{font-family:Georgia,serif;margin:24px;color:#212529;}}"
        f"table{{width:100%;border-collapse:collapse;}}</style></head>"
        f"<body>{''.join(parts)}</body></html>"
    )


def apercu_excel_html(soumission: dict) -> str:
    """
    Construit un aperçu HTML du tableau financier (rendu fidèle au Excel).

    Args:
        soumission: Données financières calculées.

    Returns:
        Fragment HTML (table) pour st.markdown().
    """
    postes   = soumission.get("postes") or soumission.get("postes_budgetaires", [])
    sous_ht  = soumission.get("sous_total_ht") or soumission.get("total_ht", 0)
    tps      = soumission.get("tps", sous_ht * 0.05)
    tvq      = soumission.get("tvq", sous_ht * 0.09975)
    ttc      = soumission.get("total_ttc", 0)
    th       = soumission.get("total_heures", 0)

    rows = "".join(
        f"<tr style='background:{'#f5f7fa' if i % 2 == 0 else '#fff'};'>"
        f"<td style='padding:5px 8px;border:1px solid #dee2e6;'>{p.get('description','')}</td>"
        f"<td style='padding:5px 8px;border:1px solid #dee2e6;text-align:center;'>"
        f"{p.get('heures_total', p.get('heures', 0))}</td>"
        f"<td style='padding:5px 8px;border:1px solid #dee2e6;text-align:center;'>"
        f"{p.get('quantite', p.get('heures', 0))}</td>"
        f"<td style='padding:5px 8px;border:1px solid #dee2e6;text-align:right;'>"
        f"${p.get('taux_horaire', p.get('taux', 0)):,.2f}</td>"
        f"<td style='padding:5px 8px;border:1px solid #dee2e6;text-align:right;font-weight:bold;'>"
        f"${p.get('total_poste', p.get('total', 0)):,.2f}</td></tr>"
        for i, p in enumerate(postes)
    )
    return f"""
<table style='width:100%;font-size:11px;border-collapse:collapse;'>
  <tr style='background:#1e3a5f;color:white;'>
    <th style='padding:6px 8px;text-align:left;'>Description</th>
    <th style='padding:6px 8px;'>H total</th>
    <th style='padding:6px 8px;'>Qté</th>
    <th style='padding:6px 8px;'>Taux $/h</th>
    <th style='padding:6px 8px;'>Total</th>
  </tr>
  {rows}
  <tr style='background:#e8f0fe;'>
    <td colspan='2' style='padding:5px 8px;border:1px solid #dee2e6;'><b>Total : {th:.0f} h</b></td>
    <td colspan='2' style='padding:5px 8px;border:1px solid #dee2e6;text-align:right;'>Sous-total HT</td>
    <td style='padding:5px 8px;border:1px solid #dee2e6;text-align:right;font-weight:bold;'>${sous_ht:,.2f}</td>
  </tr>
  <tr style='background:#e8f0fe;'>
    <td colspan='4' style='padding:5px 8px;border:1px solid #dee2e6;text-align:right;'>TPS (5 %)</td>
    <td style='padding:5px 8px;border:1px solid #dee2e6;text-align:right;'>${tps:,.2f}</td>
  </tr>
  <tr style='background:#e8f0fe;'>
    <td colspan='4' style='padding:5px 8px;border:1px solid #dee2e6;text-align:right;'>TVQ (9,975 %)</td>
    <td style='padding:5px 8px;border:1px solid #dee2e6;text-align:right;'>${tvq:,.2f}</td>
  </tr>
  <tr style='background:#1e3a5f;color:white;font-weight:bold;'>
    <td colspan='4' style='padding:6px 8px;text-align:right;'>TOTAL TTC</td>
    <td style='padding:6px 8px;text-align:right;font-size:13px;'>${ttc:,.2f}</td>
  </tr>
</table>"""


# ─────────────────────────────────────────────────────────────────────────────
# 5. AFFICHAGE TÉLÉCHARGEMENTS + APERÇUS (UI Streamlit)
# ─────────────────────────────────────────────────────────────────────────────

def afficher_telechargements(
    offre_tech: dict,
    soumission: dict,
    num_clean: str,
    user: dict,
    suffix: str = "",
) -> None:
    """
    Affiche les 3 blocs téléchargement + aperçu (Word, Excel, PDF).
    Utilisé dans l'onglet Offre Financière et l'onglet Envoi.

    Args:
        offre_tech: Mémoire technique (pour l'aperçu Word).
        soumission: Données financières (pour l'aperçu Excel).
        num_clean: Nom de projet nettoyé (pour les noms de fichiers).
        user: Profil entreprise.
        suffix: Suffixe pour les clés Streamlit (évite les doublons).
    """
    exigences  = st.session_state.offre_data.get("exigences", {})
    docx_bytes = st.session_state.get("pdf_soumission_docx") or st.session_state.get("_docx_bytes")
    xlsx_bytes = st.session_state.get("pdf_soumission_xlsx") or st.session_state.get("_xlsx_bytes")
    pdf_bytes  = st.session_state.get("pdf_soumission")      or st.session_state.get("_pdf_bytes")

    col_w, col_x, col_p = st.columns(3)

    # ── Word ─────────────────────────────────────────────────────
    with col_w:
        st.markdown("**📄 Offre Technique (Word)**")
        if docx_bytes:
            st.download_button(
                "⬇️ Télécharger le Word", data=docx_bytes,
                file_name=f"Offre_Technique_{num_clean}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"dl_docx_{suffix}", use_container_width=True
            )
            if st.button("👁️ Aperçu Word", key=f"btn_apercu_w_{suffix}", use_container_width=True):
                st.session_state[f"show_apercu_w_{suffix}"] = \
                    not st.session_state.get(f"show_apercu_w_{suffix}", False)
        else:
            st.info("Générez d'abord les documents.")

    # ── Excel ────────────────────────────────────────────────────
    with col_x:
        st.markdown("**📊 Offre Financière (Excel)**")
        if xlsx_bytes:
            st.download_button(
                "⬇️ Télécharger l'Excel", data=xlsx_bytes,
                file_name=f"Offre_Financiere_{num_clean}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key=f"dl_xlsx_{suffix}", use_container_width=True
            )
            if st.button("👁️ Aperçu Excel", key=f"btn_apercu_x_{suffix}", use_container_width=True):
                st.session_state[f"show_apercu_x_{suffix}"] = \
                    not st.session_state.get(f"show_apercu_x_{suffix}", False)
        else:
            st.info("Générez d'abord les documents.")

    # ── PDF ──────────────────────────────────────────────────────
    with col_p:
        st.markdown("**📑 Offre Complète (PDF)**")
        if pdf_bytes:
            st.download_button(
                "⬇️ Télécharger le PDF", data=pdf_bytes,
                file_name=f"Offre_Complete_{num_clean}.pdf",
                mime="application/pdf",
                key=f"dl_pdf_{suffix}", use_container_width=True
            )
            if st.button("👁️ Aperçu PDF", key=f"btn_apercu_p_{suffix}", use_container_width=True):
                st.session_state[f"show_apercu_p_{suffix}"] = \
                    not st.session_state.get(f"show_apercu_p_{suffix}", False)
        else:
            st.info("Générez d'abord les documents.")

    # ── Rendu des aperçus ────────────────────────────────────────
    import base64 as _b64

    if st.session_state.get(f"show_apercu_w_{suffix}") and docx_bytes:
        st.markdown("---")
        st.markdown("#### 📄 Aperçu — Offre Technique (Word)")
        components.html(
            apercu_word_html(offre_tech, exigences, user),
            height=750, scrolling=True
        )

    if st.session_state.get(f"show_apercu_x_{suffix}") and xlsx_bytes:
        st.markdown("---")
        st.markdown("#### 📊 Aperçu — Offre Financière (Excel)")
        components.html(
            f"<html><head><meta charset='utf-8'><style>"
            f"body{{font-family:Arial,sans-serif;margin:20px;}}"
            f"table{{width:100%;border-collapse:collapse;}}"
            f"</style></head><body>{apercu_excel_html(soumission)}</body></html>",
            height=600, scrolling=True
        )

    if st.session_state.get(f"show_apercu_p_{suffix}") and pdf_bytes:
        st.markdown("---")
        st.markdown("#### 📑 Aperçu — Offre Complète (PDF)")
        import base64 as _b64
        b64_pdf = _b64.b64encode(pdf_bytes).decode()
        # PDF.js via CDN — rendu universel, pas bloqué par Edge ni Firefox
        components.html(
            f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<style>
  * {{ margin:0; padding:0; box-sizing:border-box; }}
  body {{ background:#404040; }}
  #pdf-container {{ width:100%; height:795px; overflow-y:auto; background:#404040; padding:10px; }}
  canvas {{ display:block; margin:8px auto; box-shadow:0 2px 8px rgba(0,0,0,.5); }}
</style>
</head><body>
<div id="pdf-container"><p id="msg" style="color:white;padding:20px;font-family:Arial">
  Chargement du PDF…</p></div>
<script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.min.js"></script>
<script>
pdfjsLib.GlobalWorkerOptions.workerSrc =
  'https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.worker.min.js';

var b64 = "{b64_pdf}";
var bin = atob(b64);
var buf = new Uint8Array(bin.length);
for (var i = 0; i < bin.length; i++) {{ buf[i] = bin.charCodeAt(i); }}

var container = document.getElementById('pdf-container');

pdfjsLib.getDocument({{ data: buf }}).promise.then(function(pdf) {{
  document.getElementById('msg').remove();
  var total = pdf.numPages;
  for (var p = 1; p <= total; p++) {{
    (function(pageNum) {{
      pdf.getPage(pageNum).then(function(page) {{
        var vp = page.getViewport({{ scale: 1.5 }});
        var canvas = document.createElement('canvas');
        canvas.width  = vp.width;
        canvas.height = vp.height;
        container.appendChild(canvas);
        page.render({{ canvasContext: canvas.getContext('2d'), viewport: vp }});
      }});
    }})(p);
  }}
}}).catch(function(e) {{
  document.getElementById('msg').textContent = 'Erreur PDF : ' + e.message;
}});
</script>
</body></html>""",
            height=820, scrolling=False
        )